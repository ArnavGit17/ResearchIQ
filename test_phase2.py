# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

"""
test_phase2.py
=============
ResearchIQ Phase 2 - Automated Acceptance Test Suite

Tests:
  TEST 1 – Upload TXT
  TEST 2 – Upload DOCX
  TEST 3 – Upload PDF
  TEST 4 – Preview Page
  TEST 5 – Analytics
  TEST 6 – Delete Document
  TEST 7 – Document Management Page (search/filter/sort)
  TEST 8 – Template URL resolution (no BuildError)
"""

import io
import os
import sys
import traceback

# ── bootstrap ─────────────────────────────────────────────────────────────────
os.environ["FLASK_ENV"] = "testing"
os.environ["SECRET_KEY"] = "phase2-test-key"

from app import create_app
from app_extensions import db as _db
from models.user import User
from models.document import Document

# ── helpers ───────────────────────────────────────────────────────────────────
GREEN  = "\033[92m"
RED    = "\033[91m"
YELLOW = "\033[93m"
RESET  = "\033[0m"
BOLD   = "\033[1m"

results = []

def ok(name):
    results.append((name, True, ""))
    print(f"  {GREEN}✓ PASS{RESET}  {name}")

def fail(name, reason):
    results.append((name, False, reason))
    print(f"  {RED}✗ FAIL{RESET}  {name}")
    print(f"         {YELLOW}{reason}{RESET}")

def section(title):
    print(f"\n{BOLD}{'-'*60}{RESET}")
    print(f"{BOLD}  {title}{RESET}")
    print(f"{BOLD}{'-'*60}{RESET}")

# ── minimal TXT / DOCX / PDF bytes ────────────────────────────────────────────

def make_txt_bytes():
    text = (
        "ResearchIQ Phase 2 Test Document.\n"
        "This document contains multiple sentences. "
        "It is used to verify text extraction. "
        "Word count and sentence count should be calculated correctly.\n"
        "Paragraph two starts here. NLP analysis will follow in Phase 3."
    )
    return text.encode("utf-8")


def make_docx_bytes():
    """Create a minimal valid DOCX in memory using python-docx."""
    try:
        import docx
        from io import BytesIO
        buf = BytesIO()
        doc = docx.Document()
        doc.add_heading("ResearchIQ Test DOCX", level=1)
        doc.add_paragraph("This is the first paragraph of the test document.")
        doc.add_paragraph("Second paragraph with more content for statistics.")
        doc.save(buf)
        buf.seek(0)
        return buf.read()
    except Exception as e:
        return None, str(e)


def make_pdf_bytes():
    """Create a minimal valid PDF in memory using PyMuPDF (fitz)."""
    try:
        import fitz
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "ResearchIQ Phase 2 Test PDF.\nThis PDF tests extraction via PyMuPDF.")
        buf = doc.tobytes()
        doc.close()
        return buf
    except Exception as e:
        return None, str(e)


# ── auth helper ───────────────────────────────────────────────────────────────
def register_and_login(client):
    """Register a test user and log them in. Returns user id."""
    # Register
    rv = client.post("/auth/register", data={
        "username": "testuser",
        "email": "test@researchiq.local",
        "password": "Test1234!",
        "confirm_password": "Test1234!"
    }, follow_redirects=True)
    
    # Login
    rv = client.post("/auth/login", data={
        "email": "test@researchiq.local",
        "password": "Test1234!"
    }, follow_redirects=True)
    
    with client.session_transaction() as sess:
        user_id = sess.get("_user_id")
    return user_id


# ══════════════════════════════════════════════════════════════════════════════
#  MAIN TEST RUNNER
# ══════════════════════════════════════════════════════════════════════════════
def run_tests():
    app = create_app("testing")
    
    with app.app_context():
        _db.create_all()
        
        # Disable CSRF for tests (TestingConfig already sets WTF_CSRF_ENABLED=False)
        app.config["WTF_CSRF_ENABLED"] = False
        app.config["WTF_CSRF_CHECK_DEFAULT"] = False
        
        with app.test_client() as client:
            
            # ── Auth setup ────────────────────────────────────────────────
            section("Setup: Register & Login")
            try:
                user_id = register_and_login(client)
                ok("User registration and login")
            except Exception as e:
                fail("User registration and login", str(e))
                print(f"\n{RED}FATAL: Cannot continue without authentication.{RESET}")
                return

            # ── TEST 1: TXT Upload ─────────────────────────────────────────
            section("TEST 1: Upload TXT")
            doc_ids = {}
            try:
                txt_data = make_txt_bytes()
                rv = client.post("/upload/", data={
                    "documents": (io.BytesIO(txt_data), "test_document.txt")
                }, content_type="multipart/form-data", follow_redirects=True)
                
                if rv.status_code == 200:
                    ok("TXT upload returns HTTP 200")
                else:
                    fail("TXT upload returns HTTP 200", f"Got {rv.status_code}")
                
                doc = Document.query.filter_by(file_type="txt").first()
                if doc:
                    ok("TXT database record created")
                    doc_ids["txt"] = doc.id
                    
                    if doc.file_size > 0:
                        ok(f"TXT file size recorded ({doc.file_size} bytes)")
                    else:
                        fail("TXT file size recorded", "file_size is 0")
                    
                    if doc.word_count and doc.word_count > 0:
                        ok(f"TXT word count extracted ({doc.word_count} words)")
                    else:
                        fail("TXT word count extracted", f"word_count={doc.word_count}")
                    
                    if doc.sentence_count and doc.sentence_count > 0:
                        ok(f"TXT sentence count extracted ({doc.sentence_count} sentences)")
                    else:
                        fail("TXT sentence count extracted", f"sentence_count={doc.sentence_count}")
                    
                    if doc.character_count and doc.character_count > 0:
                        ok(f"TXT character count extracted ({doc.character_count} chars)")
                    else:
                        fail("TXT character count extracted", f"character_count={doc.character_count}")
                    
                    if doc.processing_status == "success":
                        ok("TXT processing_status = success")
                    else:
                        fail("TXT processing_status = success", f"Got: {doc.processing_status}")
                    
                    if doc.cleaned_text:
                        ok("TXT cleaned_text populated")
                    else:
                        fail("TXT cleaned_text populated", "cleaned_text is empty")
                    
                    # Verify file exists on disk
                    upload_folder = app.config["UPLOAD_FOLDER"]
                    file_path = os.path.join(upload_folder, doc.filename)
                    if os.path.isfile(file_path):
                        ok("TXT file saved to disk")
                    else:
                        fail("TXT file saved to disk", f"Not found at {file_path}")
                else:
                    fail("TXT database record created", "No Document record found in DB")
                    
            except Exception as e:
                fail("TXT Upload", f"{e}\n{traceback.format_exc()}")

            # ── TEST 2: DOCX Upload ────────────────────────────────────────
            section("TEST 2: Upload DOCX")
            try:
                docx_data = make_docx_bytes()
                if isinstance(docx_data, tuple):
                    fail("DOCX bytes creation", docx_data[1])
                else:
                    rv = client.post("/upload/", data={
                        "documents": (io.BytesIO(docx_data), "test_document.docx")
                    }, content_type="multipart/form-data", follow_redirects=True)
                    
                    if rv.status_code == 200:
                        ok("DOCX upload returns HTTP 200")
                    else:
                        fail("DOCX upload returns HTTP 200", f"Got {rv.status_code}")
                    
                    doc = Document.query.filter_by(file_type="docx").first()
                    if doc:
                        ok("DOCX database record created")
                        doc_ids["docx"] = doc.id
                        
                        if doc.word_count and doc.word_count > 0:
                            ok(f"DOCX word count extracted ({doc.word_count} words)")
                        else:
                            fail("DOCX word count extracted", f"word_count={doc.word_count}")
                        
                        if doc.processing_status == "success":
                            ok("DOCX processing_status = success")
                        else:
                            fail("DOCX processing_status = success", f"Got: {doc.processing_status}")
                    else:
                        fail("DOCX database record created", "No DOCX Document in DB")
            except Exception as e:
                fail("DOCX Upload", f"{e}\n{traceback.format_exc()}")

            # ── TEST 3: PDF Upload ─────────────────────────────────────────
            section("TEST 3: Upload PDF")
            try:
                pdf_data = make_pdf_bytes()
                if isinstance(pdf_data, tuple):
                    fail("PDF bytes creation", pdf_data[1])
                else:
                    rv = client.post("/upload/", data={
                        "documents": (io.BytesIO(pdf_data), "test_document.pdf")
                    }, content_type="multipart/form-data", follow_redirects=True)
                    
                    if rv.status_code == 200:
                        ok("PDF upload returns HTTP 200")
                    else:
                        fail("PDF upload returns HTTP 200", f"Got {rv.status_code}")
                    
                    doc = Document.query.filter_by(file_type="pdf").first()
                    if doc:
                        ok("PDF database record created")
                        doc_ids["pdf"] = doc.id
                        
                        if doc.word_count and doc.word_count > 0:
                            ok(f"PDF word count extracted ({doc.word_count} words)")
                        else:
                            fail("PDF word count extracted", f"word_count={doc.word_count}")
                        
                        if doc.processing_status == "success":
                            ok("PDF processing_status = success")
                        else:
                            fail("PDF processing_status = success", f"Got: {doc.processing_status}")
                    else:
                        fail("PDF database record created", "No PDF Document in DB")
            except Exception as e:
                fail("PDF Upload", f"{e}\n{traceback.format_exc()}")

            # ── TEST 4: Preview Page ───────────────────────────────────────
            section("TEST 4: Document Preview Page")
            for ftype, doc_id in doc_ids.items():
                try:
                    rv = client.get(f"/documents/{doc_id}", follow_redirects=True)
                    if rv.status_code == 200:
                        ok(f"Preview page renders for {ftype.upper()} (HTTP 200)")
                    else:
                        fail(f"Preview page renders for {ftype.upper()}", f"Got HTTP {rv.status_code}")
                    
                    body = rv.data.decode("utf-8", errors="replace")
                    checks = [
                        ("filename visible", "test_document"),
                        ("upload date visible", "Uploaded"),
                        # file_type is stored/rendered lowercase; CSS text-uppercase handles display
                        ("file type visible", ftype.lower()),
                        ("word count visible", "Word Count"),
                        ("status visible", "Properties"),
                    ]
                    for label, needle in checks:
                        if needle in body:
                            ok(f"Preview [{ftype}]: {label}")
                        else:
                            fail(f"Preview [{ftype}]: {label}", f"'{needle}' not found in page")
                except Exception as e:
                    fail(f"Preview page [{ftype}]", str(e))

            # ── TEST 5: Analytics ──────────────────────────────────────────
            section("TEST 5: Analytics Page")
            try:
                rv = client.get("/analytics/", follow_redirects=True)
                if rv.status_code == 200:
                    ok("Analytics page renders (HTTP 200)")
                else:
                    fail("Analytics page renders", f"Got HTTP {rv.status_code}")
                
                body = rv.data.decode("utf-8", errors="replace")
                
                if "total_words" in body or "Words" in body or "word" in body.lower():
                    ok("Analytics: word count metric visible")
                else:
                    fail("Analytics: word count metric visible", "No word count in analytics page")
                
                # Check that document count > 0 is reflected
                total_docs = Document.query.count()
                if total_docs >= 3:
                    ok(f"Analytics: {total_docs} documents tracked in DB")
                else:
                    fail("Analytics: documents tracked", f"Expected ≥3 documents, got {total_docs}")
                    
            except Exception as e:
                fail("Analytics", f"{e}\n{traceback.format_exc()}")

            # ── TEST 5b: Document Management Page ─────────────────────────
            section("TEST 5b: Document Management Page (/documents)")
            try:
                rv = client.get("/documents/", follow_redirects=True)
                if rv.status_code == 200:
                    ok("Documents management page renders (HTTP 200)")
                else:
                    fail("Documents management page renders", f"Got HTTP {rv.status_code}")
                
                body = rv.data.decode("utf-8", errors="replace")
                
                for needle in ["Search", "Sort", "Type"]:
                    if needle in body:
                        ok(f"Documents page: '{needle}' filter/control visible")
                    else:
                        fail(f"Documents page: '{needle}' visible", f"Not found in page")
                
                # Test search
                rv2 = client.get("/documents/?q=test_document", follow_redirects=True)
                if rv2.status_code == 200:
                    ok("Documents search returns HTTP 200")
                else:
                    fail("Documents search", f"Got HTTP {rv2.status_code}")
                
                # Test type filter
                rv3 = client.get("/documents/?type=txt", follow_redirects=True)
                if rv3.status_code == 200:
                    ok("Documents type filter (txt) returns HTTP 200")
                else:
                    fail("Documents type filter", f"Got HTTP {rv3.status_code}")
                    
            except Exception as e:
                fail("Document Management Page", f"{e}\n{traceback.format_exc()}")

            # ── TEST 6: Delete Document ────────────────────────────────────
            section("TEST 6: Delete Document")
            for ftype, doc_id in list(doc_ids.items()):
                try:
                    doc_before = Document.query.get(doc_id)
                    filename_before = doc_before.filename if doc_before else None
                    
                    rv = client.post(f"/upload/delete/{doc_id}", follow_redirects=True)
                    
                    if rv.status_code == 200:
                        ok(f"Delete [{ftype}] returns HTTP 200")
                    else:
                        fail(f"Delete [{ftype}] returns HTTP 200", f"Got HTTP {rv.status_code}")
                    
                    doc_after = Document.query.get(doc_id)
                    if doc_after is None:
                        ok(f"Delete [{ftype}]: DB record removed")
                    else:
                        fail(f"Delete [{ftype}]: DB record removed", "Record still in DB")
                    
                    # Verify file removed from disk
                    if filename_before:
                        upload_folder = app.config["UPLOAD_FOLDER"]
                        file_path = os.path.join(upload_folder, filename_before)
                        if not os.path.exists(file_path):
                            ok(f"Delete [{ftype}]: file removed from disk")
                        else:
                            fail(f"Delete [{ftype}]: file removed from disk", f"File still at {file_path}")
                    
                except Exception as e:
                    fail(f"Delete [{ftype}]", f"{e}\n{traceback.format_exc()}")

            # ── TEST 7: Template URL Resolution ───────────────────────────
            section("TEST 7: URL Resolution (no BuildError)")
            url_tests = [
                ("/upload/", "Upload page"),
                ("/documents/", "Documents management"),
                ("/analytics/", "Analytics"),
                ("/dashboard/", "Dashboard"),
            ]
            for url, label in url_tests:
                try:
                    rv = client.get(url, follow_redirects=True)
                    if rv.status_code == 200:
                        ok(f"URL {url} → {label}: HTTP 200")
                    elif rv.status_code in (302, 301):
                        ok(f"URL {url} → {label}: redirect (unauthenticated)")
                    else:
                        fail(f"URL {url}", f"HTTP {rv.status_code}")
                except Exception as e:
                    fail(f"URL {url}", str(e))

    # ── Summary ───────────────────────────────────────────────────────────────
    print(f"\n{BOLD}{'='*60}{RESET}")
    print(f"{BOLD}  PHASE 2 TEST RESULTS{RESET}")
    print(f"{BOLD}{'='*60}{RESET}")
    
    passed = sum(1 for _, ok_flag, _ in results if ok_flag)
    failed = sum(1 for _, ok_flag, _ in results if not ok_flag)
    total  = len(results)
    
    for name, ok_flag, reason in results:
        icon = f"{GREEN}✓{RESET}" if ok_flag else f"{RED}✗{RESET}"
        print(f"  {icon} {name}")
        if not ok_flag and reason:
            # truncate long reasons
            short = reason.split("\n")[0][:100]
            print(f"      {YELLOW}↳ {short}{RESET}")
    
    print(f"\n{BOLD}  Total: {total}  │  {GREEN}Passed: {passed}{RESET}  │  {RED}Failed: {failed}{RESET}")
    
    if failed == 0:
        print(f"\n  {GREEN}{BOLD}[OK] ALL TESTS PASSED - PHASE 2 COMPLETE{RESET}")
    else:
        print(f"\n  {RED}{BOLD}[!!] {failed} test(s) failed - see details above{RESET}")
    print()
    return failed == 0


if __name__ == "__main__":
    success = run_tests()
    sys.exit(0 if success else 1)
